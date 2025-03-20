import os
import re
import hashlib
import json
from typing import List, Dict, Any, Optional, Tuple
from fastapi import UploadFile, HTTPException
import docx
import PyPDF2
import csv
import markdown
from bs4 import BeautifulSoup
import numpy as np
import logging
import tempfile
from pathlib import Path
from pypdf import PdfReader
import pandas as pd
from sqlalchemy.orm import Session
from app.db.models.document import Document, DocumentChunk, ProcessingStatus
import asyncio
from app.services.local_embedding_service import LocalEmbeddingService
import textwrap

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Сервис для обработки документов, извлечения текста и разбиения на чанки"""
    
    # Поддерживаемые типы файлов
    SUPPORTED_FILETYPES = {
        "pdf": ["application/pdf", "pdf"],
        "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"],
        "doc": ["application/msword", "doc"],
        "txt": ["text/plain", "txt"],
        "md": ["text/markdown", "md"],
        "html": ["text/html", "html", "htm"],
        "csv": ["text/csv", "csv"],
        "json": ["application/json", "json"],
        "xlsx": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"]
    }
    
    def __init__(
        self, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        uploads_dir: Optional[str] = None,
        document_id: int = None,
        db: Session = None
    ):
        """
        Инициализация процессора документов
        
        Args:
            chunk_size: Максимальное количество символов в чанке
            chunk_overlap: Перекрытие между соседними чанками
            uploads_dir: Директория для временного хранения файлов
            document_id: ID документа
            db: Сессия базы данных
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.uploads_dir = uploads_dir or os.path.join(tempfile.gettempdir(), "document_uploads")
        
        # Создаем директорию для загрузок, если ее нет
        os.makedirs(self.uploads_dir, exist_ok=True)
        
        self.document_id = document_id
        self.db = db
        self.document = None
        if document_id and db:
            self.document = db.query(Document).filter(Document.id == document_id).first()
            if not self.document:
                logger.error(f"Документ с ID {document_id} не найден в БД")
                raise ValueError(f"Документ с ID {document_id} не найден")
            logger.info(f"Документ с ID {document_id} успешно загружен из БД")
    
    def is_filetype_supported(self, file_type: str) -> bool:
        """
        Проверяет, поддерживается ли тип файла
        
        Args:
            file_type: Тип файла или MIME-тип
            
        Returns:
            True, если тип поддерживается
        """
        file_type = file_type.lower()
        
        # Проверяем по расширению или MIME-типу
        for supported_types in self.SUPPORTED_FILETYPES.values():
            if file_type in supported_types:
                return True
        
        return False
    
    def detect_filetype(self, filename: str, content_type: Optional[str] = None) -> str:
        """
        Определяет тип файла по имени и/или MIME-типу
        
        Args:
            filename: Имя файла
            content_type: MIME-тип (опционально)
            
        Returns:
            Тип файла
        """
        # Сначала пробуем определить по MIME-типу
        if content_type:
            content_type = content_type.lower()
            for file_type, supported_types in self.SUPPORTED_FILETYPES.items():
                if content_type in supported_types:
                    return file_type
        
        # Если не удалось, определяем по расширению
        ext = Path(filename).suffix.lower().lstrip(".")
        for file_type, supported_types in self.SUPPORTED_FILETYPES.items():
            if ext in supported_types:
                return file_type
        
        # Если не удалось определить
        raise ValueError(f"Unsupported file type: {filename}, {content_type}")
    
    def calculate_content_hash(self, file_content: bytes) -> str:
        """
        Вычисляет хеш содержимого файла
        
        Args:
            file_content: Содержимое файла в байтах
            
        Returns:
            Хеш-сумма (SHA-256)
        """
        return hashlib.sha256(file_content).hexdigest()
    
    def save_file(self, file_content: bytes, filename: str) -> str:
        """
        Сохраняет файл на диск
        
        Args:
            file_content: Содержимое файла в байтах
            filename: Имя файла
            
        Returns:
            Путь к сохраненному файлу
        """
        content_hash = self.calculate_content_hash(file_content)
        
        # Создаем имя файла с хешем для уникальности
        filename_with_hash = f"{content_hash}_{os.path.basename(filename)}"
        file_path = os.path.join(self.uploads_dir, filename_with_hash)
        
        # Сохраняем файл
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        return file_path
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        content_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Обрабатывает документ: извлекает текст, метаданные и разбивает на чанки
        
        Args:
            file_content: Содержимое файла в байтах
            filename: Имя файла
            content_type: MIME-тип (опционально)
            
        Returns:
            Словарь с информацией о документе и чанках
        """
        try:
            # Определяем тип файла
            file_type = self.detect_filetype(filename, content_type)
            
            # Вычисляем хеш содержимого
            content_hash = self.calculate_content_hash(file_content)
            
            # Сохраняем файл
            file_path = self.save_file(file_content, filename)
            
            # Извлекаем текст и метаданные
            text, metadata = await self.extract_text_and_metadata(file_content, file_type, file_path)
            
            # Разбиваем текст на чанки
            chunks = self.split_text(text, metadata)
            
            # Формируем результат
            result = {
                "filename": os.path.basename(filename),
                "file_type": file_type,
                "file_size": len(file_content),
                "file_path": file_path,
                "content_hash": content_hash,
                "metadata": metadata,
                "chunks": chunks
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise
    
    async def extract_text_and_metadata(
        self, 
        file_content: bytes, 
        file_type: str,
        file_path: str
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Извлекает текст и метаданные из документа
        
        Args:
            file_content: Содержимое файла в байтах
            file_type: Тип файла
            file_path: Путь к файлу
            
        Returns:
            Кортеж (текст, метаданные)
        """
        metadata = {}
        text = ""
        
        if file_type == "pdf":
            text, metadata = self._process_pdf(file_content)
        elif file_type == "docx":
            text, metadata = self._process_docx(file_path)
        elif file_type == "txt":
            text = file_content.decode("utf-8", errors="ignore")
        elif file_type == "md":
            text = self._process_markdown(file_content)
        elif file_type == "html":
            text, metadata = self._process_html(file_content)
        elif file_type == "csv":
            text = self._process_csv(file_path)
        elif file_type == "json":
            text = self._process_json(file_content)
        elif file_type == "xlsx":
            text = self._process_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return text, metadata
    
    def _process_pdf(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Обрабатывает PDF файл"""
        metadata = {}
        text = ""
        
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                reader = PdfReader(temp_file.name)
                
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "pages": len(reader.pages)
                }
                
                pages_text = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(f"[Page {i+1}]\n{page_text}")
                
                text = "\n\n".join(pages_text)
                os.unlink(temp_file.name)
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
            
        return text, metadata
    
    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Обрабатывает DOCX файл"""
        try:
            doc = docx.Document(file_path)
            
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "paragraphs": len(doc.paragraphs)
            }
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts), metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    def _process_markdown(self, file_content: bytes) -> str:
        """Обрабатывает Markdown файл"""
        try:
            md_text = file_content.decode('utf-8', errors='ignore')
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text(separator='\n\n')
            return text
        except Exception as e:
            logger.error(f"Error processing Markdown: {str(e)}")
            raise
    
    def _process_html(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Обрабатывает HTML файл"""
        try:
            html_text = file_content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_text, 'html.parser')
            
            metadata = {
                "title": soup.title.string if soup.title else "",
                "meta_description": soup.find('meta', {'name': 'description'}).get('content', '') if soup.find('meta', {'name': 'description'}) else "",
                "h1_headers": [h1.get_text() for h1 in soup.find_all('h1')],
                "links_count": len(soup.find_all('a')),
                "images_count": len(soup.find_all('img'))
            }
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator='\n\n')
            text = re.sub(r'\n\s*\n', '\n\n', text)
            text = text.strip()
            
            return text, metadata
        except Exception as e:
            logger.error(f"Error processing HTML: {str(e)}")
            raise
    
    def _process_csv(self, file_path: str) -> str:
        """Обрабатывает CSV файл"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception:
            # Если pandas не справился, пробуем обычный CSV
            with open(file_path, "r", newline="", encoding="utf-8") as f:
                reader = csv.reader(f)
                return "\n".join([",".join(row) for row in reader])
    
    def _process_json(self, file_content: bytes) -> str:
        """Обрабатывает JSON файл"""
        try:
            data = json.loads(file_content)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON file: {str(e)}")
    
    def _process_xlsx(self, file_path: str) -> str:
        """Обрабатывает XLSX файл"""
        all_text = []
        
        # Загружаем все листы
        dfs = pd.read_excel(file_path, sheet_name=None)
        
        for sheet_name, df in dfs.items():
            all_text.append(f"[Sheet: {sheet_name}]")
            all_text.append(df.to_string(index=False))
        
        return "\n\n".join(all_text)
    
    def split_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Разбивает текст на семантически значимые чанки с учетом структуры документа.
        Размер чанка адаптируется, чтобы избежать обрезания смысловых единиц.
        """
        logger.info(f"Начинаем разбиение текста на чанки. Размер текста: {len(text)} символов")
        
        # Конфигурация разбиения
        max_chunk_size = self.chunk_size  # Используем настройку из класса
        min_chunk_size = max(100, int(max_chunk_size * 0.1))  # Минимальный размер 10% от максимального
        overlap_size = self.chunk_overlap  # Используем настройку из класса
        
        # Предобработка текста для удаления лишних пробелов
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Определяем маркеры страниц, если они есть
        page_pattern = r'(?:\[Page\s+(\d+)\]|--- Page (\d+) ---|Page (\d+):)'
        page_markers = list(re.finditer(page_pattern, text))
        
        chunks = []
        chunk_order = 0
        
        # Если есть маркеры страниц
        if page_markers:
            logger.info(f"Найдены маркеры страниц: {len(page_markers)}")
            
            # Разбиваем по страницам
            current_page = 1
            last_end = 0
            
            for i, marker in enumerate(page_markers):
                # Извлекаем номер страницы
                page_groups = marker.groups()
                page_num = next((int(pg) for pg in page_groups if pg), current_page)
                
                # Текст до следующего маркера
                start_pos = marker.end()
                end_pos = page_markers[i+1].start() if i+1 < len(page_markers) else len(text)
                
                # Если есть текст между последним маркером и текущим, его тоже обрабатываем
                if last_end < marker.start() and marker.start() - last_end > min_chunk_size:
                    unmarked_text = text[last_end:marker.start()]
                    self._process_text_segment(
                        unmarked_text, chunks, chunk_order, None, 
                        max_chunk_size, min_chunk_size, metadata
                    )
                    chunk_order += 1
                
                # Обрабатываем текст текущей страницы
                page_text = text[start_pos:end_pos]
                if page_text:
                    new_chunks = self._process_text_segment(
                        page_text, chunks, chunk_order, page_num,
                        max_chunk_size, min_chunk_size, metadata
                    )
                    chunk_order += len(new_chunks)
                
                last_end = end_pos
                current_page = page_num + 1
            
            # Обрабатываем оставшийся текст после последнего маркера
            if last_end < len(text) and len(text) - last_end > min_chunk_size:
                remaining_text = text[last_end:]
                self._process_text_segment(
                    remaining_text, chunks, chunk_order, current_page,
                    max_chunk_size, min_chunk_size, metadata
                )
        else:
            # Если нет маркеров страниц, разбиваем по семантическим единицам
            logger.info(f"Маркеры страниц не найдены, разбиваем по семантическим блокам")
            
            # Пытаемся разбить текст на логические блоки (заголовки, параграфы и т.д.)
            sections = self._split_by_sections(text)
            
            if len(sections) > 1:
                logger.info(f"Найдено {len(sections)} секций для разбиения")
                current_chunk_text = ""
                current_chunk_sections = []
                
                for section in sections:
                    # Если добавление секции превысит максимальный размер, создаем новый чанк
                    if len(current_chunk_text) + len(section) > max_chunk_size and current_chunk_text:
                        # Создаем чанк из накопленных секций
                        chunks.append({
                            "content": current_chunk_text,
                            "chunk_index": chunk_order,  # Сохраняем индекс для совместимости
                            "chunk_order": chunk_order,  # Правильное имя поля
                            "page_number": None,
                            "metadata": {
                                **metadata, 
                                "sections_count": len(current_chunk_sections),
                                "chunk_size": len(current_chunk_text)
                            }
                        })
                        
                        # Начинаем новый чанк с перекрытием
                        overlap_text = ""
                        if overlap_size > 0 and current_chunk_sections:
                            # Добавляем последние секции в перекрытие, пока не достигнем размера перекрытия
                            overlap_sections = []
                            for s in reversed(current_chunk_sections):
                                if len(overlap_text) + len(s) <= overlap_size:
                                    overlap_sections.insert(0, s)
                                    overlap_text = "".join(overlap_sections)
                                else:
                                    break
                        
                        current_chunk_text = overlap_text + section
                        current_chunk_sections = overlap_sections + [section] if overlap_text else [section]
                        chunk_order += 1
                    else:
                        # Добавляем секцию к текущему чанку
                        current_chunk_text += section
                        current_chunk_sections.append(section)
                
                # Добавляем последний чанк, если он не пустой
                if current_chunk_text and len(current_chunk_text) >= min_chunk_size:
                    chunks.append({
                        "content": current_chunk_text,
                        "chunk_index": chunk_order,  # Сохраняем индекс для совместимости
                        "chunk_order": chunk_order,  # Правильное имя поля
                        "page_number": None,
                        "metadata": {
                            **metadata, 
                            "sections_count": len(current_chunk_sections),
                            "chunk_size": len(current_chunk_text)
                        }
                    })
            else:
                # Если не удалось разбить на секции, используем более простой подход
                # разбивая текст на параграфы или предложения
                logger.info("Не удалось выделить секции, разбиваем по параграфам")
                
                # Разбиваем по параграфам
                paragraphs = [p for p in re.split(r"\n\s*\n", text) if p.strip()]
                
                # Если слишком мало параграфов, разбиваем на предложения
                if len(paragraphs) <= 1:
                    logger.info("Мало параграфов, разбиваем по предложениям")
                    paragraphs = []
                    # Попытка разбить на предложения
                    sentence_pattern = r'(?<=[.!?])\s+'
                    sentences = re.split(sentence_pattern, text)
                    for s in sentences:
                        if s.strip():
                            if len(s) > max_chunk_size:  
                                # Разбиваем длинные предложения
                                long_parts = textwrap.wrap(s, width=max_chunk_size)
                                paragraphs.extend(long_parts)
                            else:
                                paragraphs.append(s)
                
                current_chunk = ""
                current_chunk_paragraphs = []
                
                for para in paragraphs:
                    # Если текущий параграф сам по себе больше max_chunk_size, разбиваем его
                    if len(para) > max_chunk_size:
                        # Сначала добавляем то, что уже накопили
                        if current_chunk:
                            chunks.append({
                                "content": current_chunk,
                                "chunk_index": chunk_order,  # Сохраняем индекс для совместимости
                                "chunk_order": chunk_order,  # Правильное имя поля
                                "page_number": None,
                                "metadata": {**metadata, "paragraph_count": len(current_chunk_paragraphs)}
                            })
                            chunk_order += 1
                            current_chunk = ""
                            current_chunk_paragraphs = []
                        
                        # Теперь разбиваем длинный параграф
                        para_parts = textwrap.wrap(para, width=max_chunk_size)
                        for part in para_parts:
                            chunks.append({
                                "content": part,
                                "chunk_index": chunk_order,  # Сохраняем индекс для совместимости
                                "chunk_order": chunk_order,  # Правильное имя поля
                                "page_number": None,
                                "metadata": {**metadata, "paragraph_count": 1, "is_paragraph_part": True}
                            })
                            chunk_order += 1
                    # Если добавление параграфа превысит максимальный размер, создаем новый чанк
                    elif current_chunk and len(current_chunk) + len(para) + 2 > max_chunk_size:  # +2 для "\n\n"
                        chunks.append({
                            "content": current_chunk,
                            "chunk_index": chunk_order,  # Сохраняем индекс для совместимости
                            "chunk_order": chunk_order,  # Правильное имя поля
                            "page_number": None,
                            "metadata": {**metadata, "paragraph_count": len(current_chunk_paragraphs)}
                        })
                        
                        # Начинаем новый чанк с текущим параграфом
                        if self.chunk_overlap > 0 and current_chunk_paragraphs:
                            # Добавляем перекрытие
                            overlap_paras = []
                            overlap_length = 0
                            for p in reversed(current_chunk_paragraphs):
                                if overlap_length + len(p) <= self.chunk_overlap:
                                    overlap_paras.insert(0, p)
                                    overlap_length += len(p) + 2  # +2 для "\n\n"
                                else:
                                    break
                            
                            # Создаем новый чанк с перекрытием
                            current_chunk = "\n\n".join(overlap_paras + [para])
                            current_chunk_paragraphs = overlap_paras + [para]
                        else:
                            current_chunk = para
                            current_chunk_paragraphs = [para]
                        
                        chunk_order += 1
                    else:
                        # Добавляем к текущему чанку
                        if current_chunk:
                            current_chunk += "\n\n" + para
                        else:
                            current_chunk = para
                        current_chunk_paragraphs.append(para)
                
                # Добавляем последний чанк, если он не пустой
                if current_chunk and len(current_chunk) >= min_chunk_size:
                    chunks.append({
                        "content": current_chunk,
                        "chunk_index": chunk_order,  # Сохраняем индекс для совместимости
                        "chunk_order": chunk_order,  # Правильное имя поля
                        "page_number": None,
                        "metadata": {**metadata, "paragraph_count": len(current_chunk_paragraphs)}
                    })
        
        # Если никаких чанков не создано, но текст есть
        if not chunks and text.strip():
            logger.info("Не удалось разбить текст стандартными методами. Создаем один чанк")
            chunks.append({
                "content": text.strip(),
                "chunk_index": 0,  # Сохраняем индекс для совместимости
                "chunk_order": 0,  # Правильное имя поля
                "page_number": None,
                "metadata": {**metadata, "is_single_chunk": True}
            })
        
        logger.info(f"Создано {len(chunks)} чанков документа")
        return chunks
    
    def _process_text_segment(
        self, text_segment: str, chunks: List[Dict[str, Any]], 
        chunk_order: int, page_num: Optional[int], 
        max_chunk_size: int, min_chunk_size: int, 
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Обрабатывает сегмент текста, разбивая его на чанки.
        Возвращает количество созданных чанков.
        """
        segment_chunks_count = 0
        
        # Разбиваем на параграфы
        paragraphs = [p for p in re.split(r"\n\s*\n", text_segment) if p.strip()]
        
        # Если мало параграфов, возможно текст был без форматирования
        if len(paragraphs) <= 1 and len(text_segment) > max_chunk_size:
            # Пробуем разбить по предложениям
            sentence_pattern = r'(?<=[.!?])\s+'
            paragraphs = [s for s in re.split(sentence_pattern, text_segment) if s.strip()]
        
        current_chunk = ""
        current_paragraphs = []
        
        for para in paragraphs:
            para_size = len(para)
            
            # Если параграф слишком большой, разбиваем его
            if para_size > max_chunk_size:
                # Сначала сбрасываем накопленный чанк
                if current_chunk:
                    chunk_metadata = {
                        **metadata, 
                        "page": page_num,
                        "paragraph_count": len(current_paragraphs)
                    }
                    
                    chunks.append({
                        "content": current_chunk,
                        "chunk_index": chunk_order + segment_chunks_count,
                        "chunk_order": chunk_order + segment_chunks_count,
                        "page_number": page_num,
                        "metadata": chunk_metadata
                    })
                    segment_chunks_count += 1
                    current_chunk = ""
                    current_paragraphs = []
                
                # Разбиваем большой параграф на части
                for i in range(0, para_size, max_chunk_size):
                    part = para[i:min(i + max_chunk_size, para_size)]
                    
                    chunk_metadata = {
                        **metadata, 
                        "page": page_num,
                        "is_paragraph_fragment": True,
                        "fragment_index": i // max_chunk_size
                    }
                    
                    chunks.append({
                        "content": part,
                        "chunk_index": chunk_order + segment_chunks_count,
                        "chunk_order": chunk_order + segment_chunks_count,
                        "page_number": page_num,
                        "metadata": chunk_metadata
                    })
                    segment_chunks_count += 1
            
            # Если добавление параграфа превысит максимальный размер, создаем новый чанк
            elif current_chunk and len(current_chunk) + len(para) + 2 > max_chunk_size:  # +2 для "\n\n"
                chunk_metadata = {
                    **metadata, 
                    "page": page_num,
                    "paragraph_count": len(current_paragraphs)
                }
                
                chunks.append({
                    "content": current_chunk,
                    "chunk_index": chunk_order + segment_chunks_count,
                    "chunk_order": chunk_order + segment_chunks_count,
                    "page_number": page_num,
                    "metadata": chunk_metadata
                })
                segment_chunks_count += 1
                
                # Начинаем новый чанк
                current_chunk = para
                current_paragraphs = [para]
            else:
                # Добавляем параграф к текущему чанку
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
                current_paragraphs.append(para)
        
        # Добавляем последний чанк
        if current_chunk and len(current_chunk) >= min_chunk_size:
            chunk_metadata = {
                **metadata, 
                "page": page_num,
                "paragraph_count": len(current_paragraphs)
            }
            
            chunks.append({
                "content": current_chunk,
                "chunk_index": chunk_order + segment_chunks_count,
                "chunk_order": chunk_order + segment_chunks_count,
                "page_number": page_num,
                "metadata": chunk_metadata
            })
            segment_chunks_count += 1
        
        return segment_chunks_count
    
    def _split_by_sections(self, text: str) -> List[str]:
        """
        Разбивает текст на логические секции с учетом заголовков и структурных элементов.
        """
        # Компилируем шаблоны заголовков отдельно, чтобы избежать проблем с флагами
        header_patterns = [
            re.compile(r'#+\s+[^\n]+\n'),  # Markdown заголовки (# Заголовок)
            re.compile(r'^[A-ZА-Я0-9][A-ZА-Я0-9\s.,!?:;()-]{0,80}$\n', re.MULTILINE),  # Заголовки в верхнем регистре
            re.compile(r'^[A-ZА-Я][^а-яa-z\n]{5,80}$\n', re.MULTILINE),  # Еще один вариант заголовков
            re.compile(r'^(chapter|section|глава|раздел|часть)\s+[0-9IVXivx]+[^\n]*$\n', re.MULTILINE | re.IGNORECASE),  # Chapter/Section
            re.compile(r'^[^\n]{1,50}\n[=\-]{3,}$', re.MULTILINE)  # Заголовки с подчеркиванием
        ]
        
        # Ищем все заголовки в тексте
        all_headers = []
        for pattern in header_patterns:
            headers = list(pattern.finditer(text))
            all_headers.extend(headers)
        
        # Сортируем заголовки по позиции в тексте
        all_headers.sort(key=lambda x: x.start())
        
        # Если заголовки найдены, используем их для разбиения
        if all_headers:
            sections = []
            last_pos = 0
            
            for match in all_headers:
                # Добавляем текст до заголовка
                section_before = text[last_pos:match.start()].strip()
                if section_before:
                    sections.append(section_before)
                
                # Добавляем заголовок как отдельную секцию
                header = match.group(0)
                sections.append(header)
                
                last_pos = match.end()
            
            # Добавляем текст после последнего заголовка
            if last_pos < len(text):
                remaining = text[last_pos:].strip()
                if remaining:
                    sections.append(remaining)
            
            return sections
        else:
            # Если заголовки не найдены, пробуем разбить по другим структурным элементам
            structural_separators = [
                r'\n\s*\*\*\*+\s*\n',  # Разделители типа ***
                r'\n\s*---+\s*\n',     # Разделители типа ---
                r'\n\s*___+\s*\n',     # Разделители типа ___
                r'\n\s*=+\s*\n',       # Разделители типа ===
            ]
            
            # Создаем комбинированный шаблон
            combined_pattern = '|'.join(structural_separators)
            sections = re.split(combined_pattern, text)
            
            # Фильтруем пустые секции
            return [s.strip() for s in sections if s.strip()]
    
    def process(self) -> bool:
        """
        Обрабатывает документ: извлекает текст, разбивает на чанки и сохраняет в БД
        
        Returns:
            True, если обработка успешна, иначе False
        """
        try:
            if not self.document:
                logger.error("Документ не инициализирован")
                return False
            
            logger.info(f"Начинаем обработку документа ID: {self.document.id}")
            
            # Обновляем статус
            self.document.processing_status = ProcessingStatus.PROCESSING
            self.db.commit()
            
            # Проверяем, существует ли файл
            if not os.path.exists(self.document.file_path):
                error_msg = f"Файл не существует: {self.document.file_path}"
                logger.error(error_msg)
                self.document.processing_error = error_msg
                self.document.processing_status = ProcessingStatus.FAILED
                self.db.commit()
                return False
            
            # Читаем файл
            with open(self.document.file_path, "rb") as f:
                file_content = f.read()
            
            # Извлекаем текст и метаданные - используем синхронную версию вместо асинхронной
            logger.info(f"Извлекаем текст из файла типа {self.document.file_type}")
            text, metadata = self.extract_text_sync(file_content, self.document.file_type, self.document.file_path)
            
            if not text.strip():
                error_msg = "Из документа не удалось извлечь текст"
                logger.error(error_msg)
                self.document.processing_error = error_msg
                self.document.processing_status = ProcessingStatus.FAILED
                self.db.commit()
                return False
            
            # Разбиваем текст на чанки
            logger.info(f"Разбиваем текст на чанки, размер текста: {len(text)}")
            chunks = self.split_text(text, metadata)
            logger.info(f"Извлечено {len(chunks)} чанков из документа")
            
            # Создаем эмбеддинги с локальной моделью
            logger.info("Создаем эмбеддинги с локальной моделью")
            
            # Создаем папку для кэширования моделей
            cache_folder = os.path.join(os.path.dirname(self.document.file_path), ".models_cache")
            
            try:
                os.makedirs(cache_folder, exist_ok=True)
                # Здесь будет код для работы с моделями
            except Exception as e:
                logger.warning(f"Не удалось создать папку для кэширования моделей: {e}")
            
            try:
                # Инициализируем сервис эмбеддингов с моделью для русского языка
                embedding_service = LocalEmbeddingService(
                    model_name="paraphrase-multilingual-MiniLM-L12-v2",
                    cache_folder=cache_folder
                )
                
                # Генерируем эмбеддинги для чанков
                chunks_with_embeddings = embedding_service.generate_chunks_embeddings(chunks)
                logger.info(f"Созданы эмбеддинги для {len(chunks_with_embeddings)} чанков")
                
                # Сохраняем чанки с эмбеддингами в БД
                self.save_chunks(chunks_with_embeddings)
                
            except Exception as e:
                logger.exception(f"Ошибка при создании эмбеддингов: {str(e)}")
                # Если произошла ошибка, сохраняем чанки без эмбеддингов
                logger.info("Сохраняем чанки без эмбеддингов (будут работать только точные совпадения)")
                self.save_chunks(chunks)
            
            # Обновляем информацию о документе
            self.document.chunks_count = len(chunks)
            self.document.processing_status = ProcessingStatus.COMPLETED
            self.db.commit()
            
            logger.info(f"Документ ID: {self.document.id} успешно обработан")
            return True
            
        except Exception as e:
            logger.exception(f"Ошибка при обработке документа ID: {self.document_id}: {str(e)}")
            if self.document:
                self.document.processing_error = str(e)
                self.document.processing_status = ProcessingStatus.FAILED
                self.db.commit()
            return False
    
    def extract_text_sync(self, file_content: bytes, file_type: str, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Синхронная версия извлечения текста и метаданных из документа
        
        Args:
            file_content: Содержимое файла в байтах
            file_type: Тип файла
            file_path: Путь к файлу
            
        Returns:
            Кортеж (текст, метаданные)
        """
        metadata = {}
        text = ""
        
        if file_type == "pdf":
            text, metadata = self._process_pdf(file_content)
        elif file_type == "docx":
            text, metadata = self._process_docx(file_path)
        elif file_type == "txt":
            text = file_content.decode("utf-8", errors="ignore")
        elif file_type == "md":
            text = self._process_markdown(file_content)
        elif file_type == "html":
            text, metadata = self._process_html(file_content)
        elif file_type == "csv":
            text = self._process_csv(file_path)
        elif file_type == "json":
            text = self._process_json(file_content)
        elif file_type == "xlsx":
            text = self._process_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return text, metadata
    
    def save_chunks(self, chunks: List[Dict[str, Any]]) -> None:
        """
        Сохраняет чанки в базу данных.
        """
        try:
            # Удаляем существующие чанки для этого документа, если они есть
            self.db.query(DocumentChunk).filter(DocumentChunk.document_id == self.document_id).delete()
            
            # Сохраняем новые чанки
            for chunk_data in chunks:
                # Преобразуем chunk_index в chunk_order
                chunk_order = chunk_data.get("chunk_index", 0)
                
                chunk = DocumentChunk(
                    document_id=self.document_id,
                    content=chunk_data["content"],
                    chunk_order=chunk_order,  # Используем правильное имя поля
                    page_number=chunk_data.get("page_number"),
                    chunk_metadata=chunk_data.get("metadata", {})
                )
                self.db.add(chunk)
            
            self.db.commit()
            logger.info(f"Сохранено {len(chunks)} чанков в БД для документа ID: {self.document_id}")
        except Exception as e:
            logger.exception(f"Ошибка при сохранении чанков: {str(e)}")
            self.db.rollback()
            raise 